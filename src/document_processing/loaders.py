from pypdf import PdfReader
from pathlib import Path
from docx import Document
from src.utils.logger import get_logger
from src.utils.error_handler import DocumentProcessingError
from src.document_processing.preprocessor import normalize_text

logger = get_logger(__name__)

def load_multiple_files(file_input):
    logger.info(f"Loading multiple files: {len(file_input)} files")
    results = ""
    for file in file_input:
        results += loadfile(file) + "\n"
    logger.info("All files loaded")
    return results
    
def loadfile(file_input):
    logger.info(f"Loading file: {file_input}")
    result = ""
    if hasattr(file_input, 'read'):
        file_name = file_input.name
        ext = Path(file_name).suffix.lower()
        if ext == ".pdf":
            reader = PdfReader(file_input)
            pages = [page.extract_text() for page in reader.pages]
            pages = [p for p in pages if p is not None]
            result = "\n".join(pages)
        elif ext == ".txt":
            result = file_input.read().decode("utf-8")
        elif ext in [".doc", ".docx"]:
            doc = Document(file_input)
            result = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text is not None])
        else:
            logger.error(f"Unsupported file type: {ext}")
            raise DocumentProcessingError(f"Unsupported file type: {ext}")
    else:
        try:
            ext = Path(file_input).suffix.lower()
            if ext == ".pdf":
                reader = PdfReader(file_input)
                pages = [page.extract_text() for page in reader.pages]
                pages = [p for p in pages if p is not None] 
                result = "\n".join(pages)
            elif ext == ".txt":
                with open(file_input, "r") as f:
                    result = f.read()
            elif ext in [".doc", ".docx"]:
                doc = Document(file_input)
                result = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text is not None])
            else:
                logger.error(f"Unsupported file type: {ext}")
                raise DocumentProcessingError(f"Unsupported file type: {ext}")
        except FileNotFoundError:
            logger.error(f"File not found: {file_input}")
            raise DocumentProcessingError(f"File not found: {file_input}")
    logger.debug(f"Extracted {len(result)} characters")
    result = normalize_text(result)
    return result